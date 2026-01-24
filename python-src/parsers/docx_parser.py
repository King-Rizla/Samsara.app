"""
DOCX parser using python-docx.

Extracts paragraphs with style information and tables.
"""
import time
from typing import List

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from parsers.base import ParseResult, TextBlock, TableData


# Heading styles that indicate section headers
HEADING_STYLES = {
    'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4',
    'Title', 'Subtitle',
    # Common CV template styles
    'Heading1', 'Heading2', 'Heading3',
}


def parse_docx(file_path: str) -> ParseResult:
    """
    Parse DOCX file and return structured content.

    Extracts:
    - Paragraphs with style information
    - Tables as list of rows
    """
    start_time = time.perf_counter()
    warnings: List[str] = []

    try:
        doc = Document(file_path)
    except PackageNotFoundError:
        raise RuntimeError(
            "Cannot open file. The file may be corrupted, "
            "password protected, or not a valid DOCX file."
        )
    except Exception as e:
        raise RuntimeError(f"Failed to open DOCX: {str(e)}")

    # Extract paragraphs
    all_text_parts = []
    all_blocks: List[TextBlock] = []
    block_index = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Get style info
        style_name = para.style.name if para.style else None
        is_heading = style_name in HEADING_STYLES if style_name else False

        # Determine font size based on style (approximate)
        # DOCX doesn't give exact positions like PDF
        font_size = None
        if is_heading:
            font_size = 14.0  # Larger for headings
        else:
            font_size = 11.0  # Default body text

        all_text_parts.append(text)

        # Create text block (no bbox for DOCX - position is sequential)
        text_block: TextBlock = {
            "text": text,
            "bbox": (0, block_index * 20, 500, (block_index + 1) * 20),  # Approximate
            "font": style_name,
            "size": font_size,
            "page": 1  # DOCX doesn't have page info without rendering
        }
        all_blocks.append(text_block)
        block_index += 1

    # Extract tables
    all_tables: List[TableData] = []

    for table_idx, table in enumerate(doc.tables):
        table_rows = []

        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                # Get cell text, handling nested paragraphs
                cell_text = "\n".join(p.text for p in cell.paragraphs).strip()
                row_cells.append(cell_text)
            table_rows.append(row_cells)

        if table_rows:
            all_tables.append(TableData(
                rows=table_rows,
                page=1  # DOCX doesn't have page info
            ))

    if all_tables:
        warnings.append(f"Extracted {len(all_tables)} table(s)")

    # Check if document is empty
    if not all_text_parts and not all_tables:
        warnings.append("Document appears to be empty or contains only images")

    elapsed_ms = int((time.perf_counter() - start_time) * 1000)

    return ParseResult(
        raw_text="\n\n".join(all_text_parts),
        blocks=all_blocks,
        tables=all_tables,
        warnings=warnings,
        parse_time_ms=elapsed_ms,
        document_type="docx",
        page_count=1  # DOCX doesn't have page count without rendering
    )
